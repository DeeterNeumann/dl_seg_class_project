"""
generate_summary_doc.py

Generates summary of all 9 training runs
for the dual-head U-Net MoNuSAC nucleus segmentation project.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def set_cell_font(cell, size=Pt(8), bold=False, font_name="Calibri"):
    """Set font properties for all runs in a cell."""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        for run in paragraph.runs:
            run.font.size = size
            run.font.bold = bold
            run.font.name = font_name


def add_formatted_table(doc, headers, rows, col_widths=None, header_color="2F5496"):
    """Add a formatted table with header row styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"

    # data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            set_cell_font(cell, size=Pt(8))
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def main():
    doc = Document()

    # ---- page setup (landscape for wide tables) ----
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # swap width/height for landscape
    section.page_width = Cm(29.7)   # A4 landscape
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ---- default font ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # title
    title = doc.add_heading(
        "Dual-Head U-Net Training Summary", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("MoNuSAC Histopathology Nucleus Segmentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph("")  # spacer

    # architecture overview
    doc.add_heading("Architecture Overview", level=1)

    arch_items = [
        ("Encoder", "ResNet-34 (ImageNet pretrained)"),
        ("Decoder", "Shared U-Net decoder (channels: 256, 128, 64, 32, 16)"),
        ("Semantic Head", "5 classes (background, epithelial, lymphocyte, neutrophil, macrophage)"),
        ("Ternary Head", "3 classes (background, inside, boundary)"),
        ("Regularization", "SpatialDropout2d (p=0.2) on shared decoder output before both heads"),
        ("Optimizer", "AdamW"),
        ("Semantic Loss", "Soft Dice loss (multiclass)"),
        ("Ternary Loss", "Focal cross-entropy + 0.5 * Soft Dice loss"),
        ("Model Selection Metric", "Combo = 0.7 * Dice_inside_micro + 0.3 * Dice_boundary_micro"),
        ("Data Augmentation (Runs 7-9)", "HFlip, VFlip, Rot90, ColorJitter, GaussianBlur, Affine, ImageNet Normalize"),
    ]

    table = doc.add_table(rows=len(arch_items), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(arch_items):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        # Bold the key column
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        for run in table.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(21)

    doc.add_paragraph("")  # spacer

    # run-by-run results
    doc.add_heading("Run-by-Run Results", level=1)

    run_headers = [
        "Run", "LR", "\u03b3", "sem_w", "wd",
        "Bnd\nWidth", "Aug", "Drop", "ImgNet\nNorm",
        "LR Schedule", "Epochs",
        "Best\nCombo", "Best\nBnd Dice", "Key Observation",
    ]

    run_data = [
        ["1", "1e-4", "2.0", "0.7", "1e-4",
         "1px", "No", "No", "No",
         "Plateau", "200",
         "0.613", "0.24", "Baseline. Severe overfitting (train\u21920.03, val\u21921.2+)"],
        ["2", "1e-4", "1.0", "0.3", "3e-4",
         "1px", "No", "No", "No",
         "Plateau", "200",
         "0.618", "0.25", "Shifted loss weight toward ternary. Marginal gain"],
        ["3", "1e-4", "2.0", "0.3", "3e-4",
         "1px", "No", "No", "No",
         "Plateau", "200",
         "0.618", "0.25", "Added Dice loss component. 1px boundaries too thin"],
        ["4", "1e-4", "2.0", "0.3", "3e-4",
         "3px", "No", "No", "No",
         "Plateau", "200",
         "0.682", "0.44", "BREAKTHROUGH: 3px boundary dilation (+76% bnd Dice)"],
        ["5", "1e-4", "0.0", "0.3", "5e-4",
         "3px", "No", "No", "No",
         "Plateau", "200",
         "0.660", "0.42", "Disabled focal loss (\u03b3=0 = plain CE). Worse"],
        ["6", "1e-4", "1.0", "0.3", "5e-4",
         "3px", "No", "No", "No",
         "Plateau", "200",
         "~0.665", "~0.42", "Reduced \u03b3. Still worse than Run 4. Overfitting persists"],
        ["7", "1e-4", "2.0", "0.3", "3e-4",
         "3px", "Yes", "0.2", "Yes",
         "Plateau", "200",
         "~0.80\u21920.67", "~0.44",
         "Added aug+dropout+norm. Overfitting eliminated! LR collapsed too fast"],
        ["8", "1e-4", "2.0", "0.3", "3e-4",
         "3px", "Yes", "0.2", "Yes",
         "Cosine(120)", "120",
         "0.683", "0.449", "BEST MODEL. Stable cosine LR. Converged by epoch ~55"],
        ["9", "3e-4", "2.0", "0.3", "3e-4",
         "3px", "Yes", "0.2", "Yes",
         "Warmup(5)\n+Cosine(95)", "100",
         "0.678", "~0.449",
         "Higher LR + warmup. Same ceiling as Run 8. Confirms ~0.68 limit"],
    ]

    # Column widths in cm
    run_col_widths = [1.0, 1.0, 0.8, 1.0, 1.0, 1.0, 0.8, 0.8, 1.0, 2.2, 1.2, 1.2, 1.2, 12.5]

    add_formatted_table(doc, run_headers, run_data, col_widths=run_col_widths)

    doc.add_paragraph("")  # spacer

    # key milestones and lessons learned
    doc.add_heading("Key Milestones & Lessons Learned", level=1)

    milestone_headers = ["Milestone", "Run", "What Changed", "Impact"]

    milestone_data = [
        ["Baseline established", "1",
         "Dual-head U-Net with ResNet-34, focal CE + Dice loss",
         "Combo 0.613, boundary Dice 0.24. Severe overfitting (train loss\u21920.03, val loss\u21921.2+)"],
        ["Loss rebalancing", "2-3",
         "SEM_LOSS_WEIGHT 0.7\u21920.3, added Dice component, tuned \u03b3/wd",
         "Marginal gains (0.618). Problem was data representation, not loss function"],
        ["3px boundary dilation", "4",
         "Widened GT boundary masks from 1px\u21923px via morphological dilation (3x3 elliptical kernel)",
         "+76% boundary Dice (0.25\u21920.44). Single biggest improvement of all runs"],
        ["Focal loss validated", "5-6",
         "Tested \u03b3=0 (plain CE) and \u03b3=1.0",
         "Both worse than \u03b3=2.0. Focal loss confirmed essential for hard boundary pixels"],
        ["Anti-overfitting trio", "7",
         "Albumentations augmentation (HFlip, VFlip, Rot90, ColorJitter, GaussianBlur, Affine) + SpatialDropout2d(0.2) + ImageNet normalization",
         "Overfitting eliminated (train/val gap collapsed). Combo peaked at ~0.80 early but ReduceLROnPlateau killed LR prematurely"],
        ["LR schedule stabilized", "8",
         "CosineAnnealingLR(T_max=120, eta_min=1e-6) replaced ReduceLROnPlateau",
         "Best sustained combo 0.683, boundary Dice 0.449. Stable convergence by epoch ~55. BEST OVERALL MODEL"],
        ["Performance ceiling confirmed", "9",
         "3x higher LR (3e-4) + 5-epoch linear warmup + cosine decay",
         "Combo 0.678 \u2014 same ceiling as Run 8. Higher LR did not unlock better optima. Confirms ~0.68 is the architecture/data ceiling"],
    ]

    milestone_col_widths = [3.5, 1.0, 8.0, 14.2]
    add_formatted_table(
        doc, milestone_headers, milestone_data,
        col_widths=milestone_col_widths, header_color="548235"
    )

    doc.add_paragraph("")  # spacer

    # hyperparameter key
    doc.add_heading("Hyperparameter Key", level=1)

    hp_items = [
        ("\u03b3 (FOCAL_GAMMA)",
         "Focal loss exponent. 0 = plain cross-entropy; 2 = standard focal loss (focuses gradient on hard-to-classify pixels)"),
        ("sem_w (SEM_LOSS_WEIGHT)",
         "Weight on semantic head loss. Total loss = sem_w \u00d7 L_sem + 1.0 \u00d7 L_ter"),
        ("wd (weight_decay)",
         "AdamW L2 regularization coefficient"),
        ("Bnd Width",
         "Width of boundary class in ternary ground truth masks. 1px (original) vs 3px (dilated via morphological operations)"),
        ("Combo",
         "Model selection metric: 0.7 \u00d7 Dice_inside_micro + 0.3 \u00d7 Dice_boundary_micro"),
        ("Bnd Dice",
         "Ternary boundary class Dice coefficient (micro-averaged). The hardest metric to improve"),
        ("Plateau",
         "ReduceLROnPlateau(patience=10, factor=0.5) \u2014 halves LR when monitored metric stalls"),
        ("Cosine(N)",
         "CosineAnnealingLR(T_max=N, eta_min=1e-6) \u2014 smooth cosine decay over N epochs"),
        ("Warmup(N)+Cosine(M)",
         "SequentialLR: LinearLR warmup for N epochs (1e-5\u2192peak LR), then CosineAnnealing for M epochs"),
    ]

    table = doc.add_table(rows=len(hp_items), cols=2)
    table.style = "Table Grid"
    for i, (term, definition) in enumerate(hp_items):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = definition
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        for run in table.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        table.rows[i].cells[0].width = Cm(4.5)
        table.rows[i].cells[1].width = Cm(22.2)
        if i % 2 == 1:
            set_cell_shading(table.rows[i].cells[0], "F2F2F2")
            set_cell_shading(table.rows[i].cells[1], "F2F2F2")

    doc.add_paragraph("")  # spacer

    # best model summary
    doc.add_heading("Best Model: Run 8", level=1)

    p = doc.add_paragraph()
    runner = p.add_run("Run 8")
    runner.bold = True
    runner.font.size = Pt(10)
    p.add_run(
        " achieved the best sustained performance across all metrics and is the recommended "
        "model for deployment. Key final metrics:"
    ).font.size = Pt(10)

    best_metrics = [
        ("Combo (inside + boundary)", "0.683"),
        ("Boundary Dice (micro)", "0.449"),
        ("Inside Dice (micro)", "~0.778"),
        ("Semantic mIoU (macro)", "~0.760"),
        ("Semantic mIoU (micro)", "~0.845"),
        ("LR Schedule", "CosineAnnealingLR(T_max=120, eta_min=1e-6)"),
        ("Training Epochs", "120 (converged by ~55)"),
        ("Checkpoint", "Run 8 best.pt"),
    ]

    table = doc.add_table(rows=len(best_metrics), cols=2)
    table.style = "Table Grid"
    for i, (metric, value) in enumerate(best_metrics):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        for run in table.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(8)

    doc.add_paragraph("")  # spacer

    # conclusion
    doc.add_heading("Conclusion", level=1)

    conclusions = [
        (
            "The two most impactful changes across all 9 runs were: "
            "(1) widening ground truth boundary masks from 1px to 3px via morphological dilation (Run 4, +76% boundary Dice), "
            "and (2) adding the anti-overfitting trio of data augmentation, spatial dropout, and ImageNet normalization "
            "(Run 7, which eliminated the severe overfitting that plagued Runs 1-6)."
        ),
        (
            "Runs 8 and 9 converged to essentially the same performance (~0.68 combo, ~0.45 boundary Dice) despite "
            "very different learning rate strategies (1e-4 cosine vs 3e-4 warmup+cosine). This strongly indicates that "
            "0.68 combo / 0.45 boundary Dice represents the practical performance ceiling for this architecture and "
            "dataset combination."
        ),
        (
            "Breaking past this ceiling would likely require architectural changes (e.g., attention mechanisms, deeper "
            "encoder such as ResNet-50/101, larger patch sizes, or multi-scale feature fusion), additional training data, "
            "or a fundamentally different approach to boundary detection rather than further hyperparameter tuning."
        ),
        (
            "Run 7's early peak of ~0.80 combo (which could not be reproduced in Runs 8-9) is now understood to have "
            "been an artifact of ReduceLROnPlateau dynamics rather than a genuinely achievable optimum, as the higher LR "
            "exploration strategy in Run 9 failed to recapture it."
        ),
    ]

    for text in conclusions:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # ---- Save ----
    out_path = Path(__file__).resolve().parent / "training_run_summary.docx"
    doc.save(str(out_path))
    print(f"Document saved to: {out_path}")


if __name__ == "__main__":
    main()
